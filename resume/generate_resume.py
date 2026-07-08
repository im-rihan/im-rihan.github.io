"""
Generate resume PDF and Word from resume.html (same visual design).
Run: python generate_resume.py
Output: docs/ and public/ — resume.html, resume.pdf, resume.docx

Word (.docx) embeds a high-DPI raster of the PDF so it matches HTML/PDF exactly.
"""

import re
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parent
HTML_SOURCE = SCRIPT_DIR / "resume.html"
DOCS_DIR = REPO_ROOT / "docs"
PUBLIC_DIR = REPO_ROOT / "public"
OUTPUT_HTML = DOCS_DIR / "resume.html"
OUTPUT_PDF = DOCS_DIR / "resume.pdf"
OUTPUT_DOCX = DOCS_DIR / "resume.docx"
OUTPUT_PNG = DOCS_DIR / "resume-page.png"
RESUME_TS = REPO_ROOT / "src" / "lib" / "resume.ts"

A4_HEIGHT_PX = 1123  # ~297mm at 96dpi
DOCX_RASTER_DPI = 220


def _resume_date_from_git() -> str:
    """Return YYYYMMDD of the last git commit that touched resume.html, or today."""
    try:
        out = subprocess.check_output(
            ["git", "log", "-1", "--format=%cd", "--date=format:%Y%m%d", str(HTML_SOURCE)],
            cwd=REPO_ROOT,
            text=True,
            stderr=subprocess.DEVNULL,
        ).strip()
        return out if out else datetime.now(timezone.utc).strftime("%Y%m%d")
    except Exception:
        return datetime.now(timezone.utc).strftime("%Y%m%d")


def update_resume_version() -> None:
    """Patch RESUME_VERSION in src/lib/resume.ts to match last git commit date."""
    if not RESUME_TS.exists():
        return
    version = _resume_date_from_git()
    content = RESUME_TS.read_text(encoding="utf-8")
    patched = re.sub(
        r'(const RESUME_VERSION = ")[^"]+(")',
        rf'\g<1>{version}\g<2>',
        content,
    )
    if patched != content:
        RESUME_TS.write_text(patched, encoding="utf-8")
        print(f"[OK] RESUME_VERSION updated → {version}")


def sync_html():
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    PUBLIC_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(HTML_SOURCE, OUTPUT_HTML)
    print(f"[OK] HTML synced: {OUTPUT_HTML}")


def sync_to_public():
    """Copy resume exports to Next.js public/ for static hosting."""
    for name in ("resume.html", "resume.pdf", "resume.docx"):
        src = DOCS_DIR / name
        if src.exists():
            shutil.copy2(src, PUBLIC_DIR / name)
            print(f"[OK] Public sync: {PUBLIC_DIR / name}")


def _html_uri() -> str:
    path = OUTPUT_HTML if OUTPUT_HTML.exists() else HTML_SOURCE
    return path.as_uri()


def _compute_pdf_scale(page) -> float:
    """Scale content to fit a single A4 page when print CSS alone is not enough."""
    metrics = page.evaluate(
        """() => {
          const root = document.querySelector('.page');
          if (!root) return { height: 0 };
          const prev = root.style.transform;
          root.style.transform = 'none';
          const height = root.getBoundingClientRect().height;
          root.style.transform = prev;
          return { height };
        }"""
    )
    content_height = metrics.get("height", 0)
    if content_height <= A4_HEIGHT_PX:
        return 1.0
    return round(min(0.95, A4_HEIGHT_PX / content_height), 3)


def _pdf_is_fresh() -> bool:
    if not OUTPUT_PDF.exists() or not OUTPUT_HTML.exists():
        return False
    return OUTPUT_PDF.stat().st_mtime >= OUTPUT_HTML.stat().st_mtime


def _ensure_playwright_chromium() -> None:
    from playwright.sync_api import sync_playwright

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch()
            browser.close()
        return
    except Exception as exc:
        message = str(exc)
        if "Executable doesn't exist" not in message and "BrowserType.launch" not in message:
            raise

    print("[INFO] Playwright Chromium not found — downloading (one-time setup)...")
    subprocess.run(
        [sys.executable, "-m", "playwright", "install", "chromium"],
        check=True,
    )
    print("[OK] Playwright Chromium installed")


def generate_pdf():
    from playwright.sync_api import sync_playwright

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    _ensure_playwright_chromium()

    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page(viewport={"width": 900, "height": 1200})
        page.emulate_media(media="print")
        page.goto(_html_uri(), wait_until="networkidle")

        scale = _compute_pdf_scale(page)
        for _ in range(6):
            page.pdf(
                path=str(OUTPUT_PDF),
                format="A4",
                print_background=True,
                margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
                prefer_css_page_size=True,
                scale=scale,
            )
            pages = pdf_page_count()
            if pages == 1:
                break
            scale = round(scale * 0.92, 3)

        browser.close()
    print(f"[OK] PDF saved: {OUTPUT_PDF} (scale={scale})")
    return scale


def pdf_page_count() -> int | None:
    try:
        from pypdf import PdfReader

        return len(PdfReader(str(OUTPUT_PDF)).pages)
    except ImportError:
        return None


def _pdf_to_png(pdf_path: Path, png_path: Path, dpi: int = DOCX_RASTER_DPI) -> None:
    import fitz

    doc = fitz.open(str(pdf_path))
    if len(doc) == 0:
        doc.close()
        raise ValueError("PDF has no pages")
    page = doc[0]
    zoom = dpi / 72
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    pix.save(str(png_path))
    doc.close()


def generate_docx(*, require_fresh_pdf: bool = False):
    """Embed a raster of the PDF — pixel-identical to the HTML/PDF resume."""
    from docx import Document
    from docx.shared import Mm, Pt

    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    if not OUTPUT_PDF.exists():
        print("[INFO] PDF missing — generating PDF first for ditto Word export...")
        generate_pdf()
    elif require_fresh_pdf and not _pdf_is_fresh():
        print("[INFO] PDF is older than HTML — regenerating PDF for ditto Word export...")
        generate_pdf()

    if not OUTPUT_PDF.exists():
        raise FileNotFoundError(
            "PDF was not generated. Run: python -m playwright install chromium"
        )

    _pdf_to_png(OUTPUT_PDF, OUTPUT_PNG, dpi=DOCX_RASTER_DPI)

    doc = Document()
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run()
    run.add_picture(str(OUTPUT_PNG), width=Mm(210), height=Mm(297))

    try:
        doc.save(str(OUTPUT_DOCX))
        print(f"[OK] Word document saved: {OUTPUT_DOCX} (mirrors PDF at {DOCX_RASTER_DPI} DPI)")
    except PermissionError:
        fallback = SCRIPT_DIR / "resume.docx"
        doc.save(str(fallback))
        print(f"[WARN] Could not overwrite {OUTPUT_DOCX} (file open in Word).")
        print(f"[OK] Word document saved: {fallback} (mirrors PDF at {DOCX_RASTER_DPI} DPI)")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Generate resume exports")
    parser.add_argument("--docx-only", action="store_true", help="Regenerate Word document only")
    parser.add_argument("--skip-pdf", action="store_true", help="Skip PDF generation")
    args = parser.parse_args()

    print("Generating resume from resume.html...\n")

    if args.docx_only:
        if not OUTPUT_HTML.exists():
            sync_html()
        generate_docx()
        sync_to_public()
        print(f"\nDone! Files in: {DOCS_DIR} and {PUBLIC_DIR}")
        raise SystemExit(0)

    sync_html()
    pdf_ok = False
    if not args.skip_pdf:
        try:
            generate_pdf()
            pdf_ok = True
            pages = pdf_page_count()
            if pages is not None:
                status = "single page" if pages == 1 else f"{pages} pages"
                print(f"[INFO] PDF page count: {pages} ({status})")
        except Exception as exc:
            print(f"[ERROR] PDF generation failed: {exc!s}")
            print("       Fix: pip install -r resume/requirements.txt")
            print("            python -m playwright install chromium")
            print("       Or run: npm run setup:resume")
    try:
        if pdf_ok or _pdf_is_fresh():
            generate_docx(require_fresh_pdf=not pdf_ok)
        else:
            print(
                "[WARN] Word skipped — PDF is missing or stale after a failed PDF run."
            )
    except Exception as exc:
        print(f"[WARN] Word generation skipped: {exc!s}")
    sync_to_public()
    update_resume_version()
    print(f"\nDone! Files in: {DOCS_DIR} and {PUBLIC_DIR}")
